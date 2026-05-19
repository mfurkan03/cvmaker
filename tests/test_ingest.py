import pytest
from io import BytesIO


def test_extract_text_from_txt():
    from app.ingest import extract_text
    content = b"I am a software engineer with 5 years experience."
    result = extract_text(content, "resume.txt")
    assert "software engineer" in result


def test_extract_text_from_txt_utf8_turkish():
    from app.ingest import extract_text
    text = "Yazılım mühendisiyim. Şirketimde çalışıyorum."
    result = extract_text(text.encode("utf-8"), "cv.txt")
    assert "Yazılım" in result
    assert "çalışıyorum" in result


def test_extract_text_from_txt_cp1254_turkish():
    from app.ingest import extract_text
    text = "İstanbul'da yazılım geliştiriyorum. Şeker gibi kod yazarım."
    result = extract_text(text.encode("cp1254"), "cv.txt")
    assert "İstanbul" in result
    assert "Şeker" in result


def test_extract_text_from_docx():
    import docx
    from app.ingest import extract_text
    doc = docx.Document()
    doc.add_paragraph("My name is Test User.")
    doc.add_paragraph("I work at Test Corp.")
    buf = BytesIO()
    doc.save(buf)
    result = extract_text(buf.getvalue(), "resume.docx")
    assert "Test User" in result
    assert "Test Corp" in result


def test_extract_text_from_pdf():
    import fitz
    from app.ingest import extract_text
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Software Engineer with Python experience.")
    pdf_bytes = doc.tobytes()
    doc.close()
    result = extract_text(pdf_bytes, "resume.pdf")
    assert "Software Engineer" in result


def test_extract_text_unsupported_raises():
    from app.ingest import extract_text
    with pytest.raises(ValueError, match="Unsupported file type"):
        extract_text(b"data", "file.xlsx")


def test_extract_text_legacy_doc_raises():
    from app.ingest import extract_text
    with pytest.raises(ValueError, match="Unsupported file type"):
        extract_text(b"data", "file.doc")


def test_extract_text_empty_raises():
    from app.ingest import extract_text
    with pytest.raises(ValueError, match="Empty file"):
        extract_text(b"", "resume.txt")
